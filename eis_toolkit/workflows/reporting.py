"""Report writers for the mineral prospectivity workflow.

Produces the three deliverables: a prospectivity map (PDF), a statistics
workbook (XLSX) and a summary report (DOCX).
"""

import os

import geopandas as gpd
import numpy as np
from beartype import beartype
from beartype.typing import Union


# Substring (case-insensitive) in a fault's type/feature attribute that marks
# it as a thrust / layered (low-angle) fault, drawn with a dashed line.
THRUST_KEYWORDS = ("thrust",)
# Attribute columns searched for the fault type description.
FAULT_TYPE_COLUMNS = ("TYPE", "FEATURE", "FAULT_TYPE", "type")


def _is_thrust(values) -> np.ndarray:
    """Boolean mask of features whose type text contains a thrust keyword."""
    text = values.astype(str).str.lower()
    mask = np.zeros(len(values), dtype=bool)
    for keyword in THRUST_KEYWORDS:
        mask |= text.str.contains(keyword, na=False)
    return mask


def _split_faults(faults: gpd.GeoDataFrame):
    """Split faults into (thrust, other) GeoDataFrames using the type attribute."""
    type_column = next((c for c in FAULT_TYPE_COLUMNS if c in faults.columns), None)
    if type_column is None:
        return faults.iloc[0:0], faults
    thrust_mask = _is_thrust(faults[type_column])
    return faults[thrust_mask], faults[~thrust_mask]


def _fault_lines(faults: gpd.GeoDataFrame) -> gpd.GeoDataFrame:
    """Return drawable line geometries for faults, dropping empty/null ones.

    Polygon fault geometries are converted to their boundaries so they read as
    lines; line geometries are kept as-is. This avoids calling ``.boundary`` on
    geometries (e.g. lines or empties) where it would raise.
    """
    cleaned = faults[~(faults.geometry.is_empty | faults.geometry.isna())]
    if cleaned.empty:
        return cleaned
    is_polygon = cleaned.geometry.type.isin(["Polygon", "MultiPolygon"])
    lines = cleaned.geometry.copy()
    lines.loc[is_polygon] = cleaned.geometry[is_polygon].boundary
    result = cleaned.set_geometry(lines)
    return result[~(result.geometry.is_empty | result.geometry.isna())]


@beartype
def write_map_pdf(
    prospectivity: np.ndarray,
    profile: dict,
    deposits: gpd.GeoDataFrame,
    output_path: Union[str, os.PathLike],
    faults: gpd.GeoDataFrame = None,
) -> None:
    """Render the prospectivity map with deposits and faults to a PDF.

    Deposits are drawn as yellow points. Faults are drawn in red: regular faults
    as solid lines and thrust (layered/low-angle) faults as dashed lines.
    """
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.lines import Line2D
    from rasterio.transform import array_bounds

    masked = np.ma.masked_equal(prospectivity, profile.get("nodata", -9999.0))
    height, width = prospectivity.shape
    left, bottom, right, top = array_bounds(height, width, profile["transform"])

    fig, ax = plt.subplots(figsize=(8.27, 11.69))  # A4 portrait
    image = ax.imshow(masked, cmap="viridis", extent=(left, right, bottom, top))
    fig.colorbar(image, ax=ax, shrink=0.6, label="Prospectivity (probability)")

    legend_handles = []

    # Faults in red: thrust faults dashed, other faults solid.
    if faults is not None and not faults.empty:
        thrust_faults, other_faults = _split_faults(faults)
        other_lines = _fault_lines(other_faults)
        thrust_lines = _fault_lines(thrust_faults)
        if not other_lines.empty:
            other_lines.plot(ax=ax, color="red", linewidth=0.8, linestyle="solid")
            legend_handles.append(Line2D([0], [0], color="red", lw=1.2, linestyle="solid", label="Fault"))
        if not thrust_lines.empty:
            thrust_lines.plot(ax=ax, color="red", linewidth=1.0, linestyle="dashed")
            legend_handles.append(Line2D([0], [0], color="red", lw=1.2, linestyle="dashed", label="Thrust fault"))

    # Deposits as yellow points.
    if not deposits.empty:
        deposit_points = deposits[deposits.geometry.type.isin(["Point", "MultiPoint"])]
        if not deposit_points.empty:
            deposit_points.plot(ax=ax, color="yellow", markersize=12, edgecolor="black", linewidth=0.3)
            legend_handles.append(
                Line2D([0], [0], marker="o", color="none", markerfacecolor="yellow",
                       markeredgecolor="black", markersize=8, label="Deposits")
            )

    if legend_handles:
        ax.legend(handles=legend_handles, loc="upper right")

    ax.set_title("Mineral Prospectivity Map")
    ax.set_xlabel("Easting")
    ax.set_ylabel("Northing")
    fig.tight_layout()
    fig.savefig(output_path, format="pdf")
    plt.close(fig)


@beartype
def write_statistics_xlsx(stats: dict, output_path: Union[str, os.PathLike]) -> None:
    """Write the statistics dictionary to an Excel workbook with three sheets."""
    from openpyxl import Workbook

    workbook = Workbook()

    summary_sheet = workbook.active
    summary_sheet.title = "Summary"
    summary_sheet.append(["Metric", "Value"])
    for key, value in stats["summary"].items():
        summary_sheet.append([key, value])

    metrics_sheet = workbook.create_sheet("Model metrics")
    metrics_sheet.append(["Metric", "Value"])
    for key, value in stats["metrics"].items():
        metrics_sheet.append([key, round(value, 4)])

    importance_sheet = workbook.create_sheet("Feature importance")
    importance_sheet.append(["Feature", "Importance"])
    for key, value in stats["feature_importance"].items():
        importance_sheet.append([key, round(value, 4)])

    relation = stats.get("fault_relation")
    if relation:
        fault_sheet = workbook.create_sheet("Fault-deposit relation")
        fault_sheet.append(["Metric", "Value"])
        fault_sheet.append([f"Buffer distance (km)", relation["buffer_km"]])
        fault_sheet.append(["Total deposits", relation["total_deposits"]])
        fault_sheet.append(["Deposits within buffer (fault-related)", relation["within_count"]])
        fault_sheet.append(["Fault-related deposits (%)", relation["within_percent"]])
        fault_sheet.append(["Deposits outside buffer", relation["outside_count"]])
        fault_sheet.append(["Non-fault-related deposits (%)", relation["outside_percent"]])

    if stats.get("comparison"):
        comparison_sheet = workbook.create_sheet("Model comparison")
        columns = ["model"] + [k for k in stats["comparison"][0] if k != "model"]
        comparison_sheet.append(columns)
        for row in stats["comparison"]:
            comparison_sheet.append(
                [row["model"]] + [round(float(row[c]), 4) for c in columns if c != "model"]
            )

    workbook.save(output_path)


@beartype
def write_report_docx(
    stats: dict,
    map_path: Union[str, os.PathLike],
    output_path: Union[str, os.PathLike],
    inputs: dict,
) -> None:
    """Write a summary Word report describing inputs, model and results."""
    from docx import Document

    document = Document()
    document.add_heading("Mineral Prospectivity Mapping Report", level=0)

    document.add_heading("Input data", level=1)
    input_table = document.add_table(rows=1, cols=2)
    input_table.style = "Light Grid Accent 1"
    header = input_table.rows[0].cells
    header[0].text = "Layer"
    header[1].text = "Source"
    for name, path in inputs.items():
        row = input_table.add_row().cells
        row[0].text = name
        row[1].text = str(path)

    document.add_heading("Method", level=1)
    document.add_paragraph(
        "Evidence layers (DEM, distance-to-fault and geology) were aligned to the "
        "DEM grid. A Random Forest classifier was trained using deposit points as "
        "positive samples and randomly sampled background cells as negatives. The "
        "trained model was applied to every cell to produce a prospectivity "
        "probability surface."
    )

    document.add_heading("Summary statistics", level=1)
    _add_dict_table(document, stats["summary"])

    document.add_heading("Model performance", level=1)
    _add_dict_table(document, {k: round(v, 4) for k, v in stats["metrics"].items()})

    if stats["feature_importance"]:
        document.add_heading("Feature importance", level=1)
        _add_dict_table(
            document, {k: round(v, 4) for k, v in stats["feature_importance"].items()}
        )

    relation = stats.get("fault_relation")
    if relation:
        document.add_heading("Fault-deposit relationship", level=1)
        document.add_paragraph(
            f"Using a {relation['buffer_km']} km buffer around faults, "
            f"{relation['within_count']} of {relation['total_deposits']} deposits "
            f"({relation['within_percent']}%) are fault-related (within the buffer), and "
            f"{relation['outside_count']} ({relation['outside_percent']}%) lie outside it."
        )
        _add_dict_table(document, {
            "Buffer distance (km)": relation["buffer_km"],
            "Total deposits": relation["total_deposits"],
            "Fault-related (within buffer)": relation["within_count"],
            "Fault-related (%)": relation["within_percent"],
            "Outside buffer": relation["outside_count"],
            "Outside buffer (%)": relation["outside_percent"],
        })

    if stats.get("comparison"):
        document.add_heading("Model comparison", level=1)
        _add_comparison_table(document, stats["comparison"])

    document.add_heading("Prospectivity map", level=1)
    document.add_paragraph(f"See accompanying map: {os.path.basename(str(map_path))}")

    document.save(output_path)


def _add_dict_table(document, data: dict) -> None:
    """Add a two-column key/value table to a Word document."""
    table = document.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "Item"
    header[1].text = "Value"
    for key, value in data.items():
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)


def _add_comparison_table(document, comparison: list) -> None:
    """Add a model comparison table (one row per model, metric columns)."""
    columns = ["model"] + [k for k in comparison[0] if k != "model"]
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    for index, name in enumerate(columns):
        table.rows[0].cells[index].text = name
    for entry in comparison:
        cells = table.add_row().cells
        for index, name in enumerate(columns):
            value = entry[name]
            cells[index].text = value if name == "model" else f"{float(value):.4f}"
